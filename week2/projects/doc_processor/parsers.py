"""Document parsers for PDF, DOCX, TXT"""

from pathlib import Path


def parse_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def parse_pdf(file_path: str) -> str:
    try:
        import fitz
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except ImportError:
        return "PyMuPDF not installed. Run: pip install pymupdf"
    except Exception as e:
        return f"PDF parsing error: {str(e)}"


def parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        return "python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"DOCX parsing error: {str(e)}"


def parse_document(file_path: str) -> tuple[str, str]:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".txt":
        return parse_txt(file_path), "txt"
    elif ext == ".pdf":
        return parse_pdf(file_path), "pdf"
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_path), "docx"
    else:
        return f"Unsupported file type: {ext}", "unknown"